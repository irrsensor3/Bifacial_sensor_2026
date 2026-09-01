import streamlit as st
import numpy as np
import pandas as pd
from docx import Document
from io import BytesIO
from datetime import datetime
import math
import os
import base64
import hashlib
import hmac
import time
import matplotlib.pyplot as plt
from fpdf import FPDF
import tempfile
from supabase import create_client

# Create a Supabase client here so other modules can import it from this package.
# Missing secrets used to raise a bare KeyError at import time, which Streamlit
# shows as a stack trace with no indication of what to fix.
def _require_secret(name: str) -> str:
    try:
        return st.secrets[name]
    except (KeyError, FileNotFoundError):
        st.error(
            f"Missing `{name}`. Add it to `.streamlit/secrets.toml` locally, or "
            f"to the app's secrets in Streamlit Cloud, then reload."
        )
        st.stop()


supabase = create_client(_require_secret("SUPABASE_URL"),
                         _require_secret("SUPABASE_KEY"))

# =========================
# AUTHENTICATION
# =========================

# The password lives in secrets, never in the repo. Previously both the salt
# and the literal password ("admin123") were in this file, which is public --
# anyone reading it had admin, and admin can reboot and shut down the Pi.
#
# Generate the hash once and put it in secrets.toml:
#     python -c "import hashlib,secrets; s=secrets.token_hex(16); \
#       print('ADMIN_SALT =', repr(s)); \
#       print('ADMIN_PASSWORD_HASH =', repr(hashlib.pbkdf2_hmac('sha256', \
#       b'YOUR-PASSWORD', s.encode(), 200_000).hex()))"
_PBKDF2_ROUNDS = 200_000
_LOCKOUT_AFTER = 5
_LOCKOUT_SECONDS = 300

@st.cache_data(ttl=5)
def fetch_system_logs(limit=200):
    rows = _safe_query("system_logs", limit, "system logs")
    if not rows:
        return pd.DataFrame()
    df = pd.DataFrame(rows)
    df["created_at"] = pd.to_datetime(df["created_at"], errors="coerce")
    return df


def system_log_panel(limit=200, height=260):
    """System log viewer. Call this once per page load — since app.py's
    render_overview() runs on every page (st.navigation reruns the whole
    script on page switch), placing the call there shows this on every page,
    right under the live array diagram."""
    df_logs = fetch_system_logs(limit=limit)
    with st.expander("System logs", expanded=False):
        if df_logs.empty:
            st.caption("No log entries yet.")
            return
        show_df = (
            df_logs.sort_values("created_at", ascending=False)[["created_at", "message"]]
        )
        st.dataframe(show_df, use_container_width=True, hide_index=True, height=height)
        
def hash_password(password: str, salt: str) -> str:
    """PBKDF2 rather than a single SHA-256 pass. A plain hash of a short
    password is brute-forced in seconds on a laptop; 200k rounds makes each
    guess cost real time."""
    return hashlib.pbkdf2_hmac(
        "sha256", password.encode(), salt.encode(), _PBKDF2_ROUNDS
    ).hex()


def check_password(password: str) -> bool:
    try:
        salt = st.secrets["ADMIN_SALT"]
        expected = st.secrets["ADMIN_PASSWORD_HASH"]
    except (KeyError, FileNotFoundError):
        st.error(
            "Admin login is not configured. Add `ADMIN_SALT` and "
            "`ADMIN_PASSWORD_HASH` to the app secrets — see the comment in "
            "ui_sections.py for the one-liner that generates them."
        )
        return False
    # constant-time compare, so response timing does not leak how much of the
    # hash matched
    return hmac.compare_digest(hash_password(password, salt), expected)


def _login_locked() -> int:
    """Seconds remaining on the lockout, 0 if not locked."""
    fails = st.session_state.get("_login_fails", 0)
    if fails < _LOCKOUT_AFTER:
        return 0
    since = time.time() - st.session_state.get("_login_last_fail", 0)
    return max(0, int(_LOCKOUT_SECONDS - since))


def login():
    """Render the login UI and set session state on success."""
    st.title("Bifacial PV logging system")
    st.caption(
        "Sign in as a guest to view live data and reports, or as an admin to "
        "also control the logger."
    )

    col1, col2 = st.columns(2)

    with col1:
        with st.container(border=True):
            st.subheader("Admin")
            st.caption("Full access, including power controls and sensor settings.")
            locked = _login_locked()
            password = st.text_input(
                "Admin password", type="password", key="admin_pass",
                disabled=locked > 0,
            )
            if st.button("Sign in as admin", disabled=locked > 0,
                         use_container_width=True):
                if check_password(password):
                    st.session_state.auth = True
                    st.session_state.user_role = "admin"
                    st.session_state._login_fails = 0
                    st.rerun()
                else:
                    # Throttle guessing. Without this the password can be
                    # attacked at the speed of the network.
                    st.session_state._login_fails = st.session_state.get("_login_fails", 0) + 1
                    st.session_state._login_last_fail = time.time()
                    left = _LOCKOUT_AFTER - st.session_state._login_fails
                    if left > 0:
                        st.error(f"That password did not match. {left} attempt(s) left.")
                    else:
                        st.error("Too many attempts. Try again in 5 minutes.")
            if locked:
                st.warning(f"Locked for {locked // 60}m {locked % 60}s after too many attempts.")

    with col2:
        with st.container(border=True):
            st.subheader("Guest")
            st.caption("View live readings, charts and reports. No password needed.")
            if st.button("Continue as guest", use_container_width=True):
                st.session_state.auth = True
                st.session_state.user_role = "guest"
                st.rerun()


def require_login(role: str | None = None):
    """Page guard for the multipage app. Call this at the top of every
    page in app_pages/. app.py already gates access to st.navigation
    behind login (and hides the admin page from non-admins), so this
    is mainly a safety net for direct reruns. If `role` is given (e.g.
    "admin"), also requires that exact role and stops with an info
    message otherwise."""
    if not st.session_state.get("auth"):
        st.warning("Please log in first.")
        st.stop()

    if role is not None and st.session_state.get("user_role") != role:
        st.info(f"ℹ️ This page is only available in {role} mode.")
        st.stop()


def inject_theme():
    """Injects the 'Solar Admin' visual theme — light background,
    white rounded cards with soft shadows, blue/green/amber accent
    colors, clean sans-serif type. Called once from app.py; since
    st.navigation reruns app.py's whole script on every page switch,
    one call there covers every page, no per-page injection needed."""
    st.markdown(
        """
        <style>
        /* @import must be the first thing inside the block, and the block must
           contain nothing but CSS. An earlier version put <link> tags above
           this: Streamlit's sanitiser rejects <link>, which invalidated the
           whole thing and printed the stylesheet on the page as plain text. */
        @import url('https://fonts.googleapis.com/css2?family=Manrope:wght@400;500;600;700;800&family=IBM+Plex+Mono:wght@400;500;600&display=swap');

        :root {
            /* The rig measures light on two sides: direct beam on the front,
               ground-reflected on the rear. Warm for direct, cool for
               reflected -- and amber against teal stays distinguishable for
               colour-blind readers, where red against green would not. */
            --sun:       #F2A03D;
            --sun-deep:  #D97706;
            --sun-soft:  #FFF6E9;
            --rear:      #0E7C86;
            --rear-soft: #E6F4F5;

            --ink:     #0F1B2A;
            --ink-mid: #56656F;
            --base:    #F5F7FA;
            --surface: #FFFFFF;
            --line:    #E3E9EF;
            --power:   #2E7D5B;
            --alert:   #C0392B;

            --sans: 'Manrope', -apple-system, BlinkMacSystemFont, 'Segoe UI',
                    Roboto, Arial, sans-serif;
            --mono: 'IBM Plex Mono', ui-monospace, Consolas, monospace;

            --r: 16px;
            --shadow: 0 2px 4px rgba(15,27,42,.04), 0 12px 28px rgba(15,27,42,.06);
        }

        .stApp { background-color: var(--base); color: var(--ink); }
        body, p, div, span, label, button, input { font-family: var(--sans); }
        
         /* Streamlit renders built-in icons (sidebar collapse arrows, etc.) as
           text glyphs that depend on the Material Symbols font. The rule
           above overrides that font on every span, including these, so the
           icon name prints as literal text instead of a glyph. Carve icons
           back out. */
        [data-testid="stIconMaterial"],
        span[class*="material-icons"],
        span[class*="material-symbols"] {
            font-family: 'Material Symbols Outlined', 'Material Symbols Rounded',
                          'Material Icons' !important;
        }

        h1 {
            font-family: var(--sans) !important;
            font-weight: 800 !important;
            font-size: 2.6rem !important;
            letter-spacing: -0.035em !important;
            color: var(--ink) !important;
        }
        h2 {
            font-weight: 700 !important;
            letter-spacing: -0.02em !important;
            color: var(--ink) !important;
        }
        h3 {
            font-weight: 700 !important;
            font-size: 1.15rem !important;
            letter-spacing: -0.01em !important;
            color: var(--ink) !important;
        }

        /* ---------- hero ---------- */
        .hero {
            position: relative;
            border-radius: var(--r);
            overflow: hidden;
            padding: 2.9rem 2.4rem 2.4rem;
            min-height: 300px;
            margin-bottom: 1.4rem;
            background-size: cover;
            background-position: center 78%;
            box-shadow: var(--shadow);
        }
        .hero::before {
            content: "";
            position: absolute; inset: 0;
            background: linear-gradient(105deg,
                        rgba(11,21,33,.96) 0%,
                        rgba(11,21,33,.88) 42%,
                        rgba(11,21,33,.55) 78%,
                        rgba(11,21,33,.42) 100%);
        }
        .hero > * { position: relative; }
        .hero-eyebrow {
            font-family: var(--mono);
            font-size: .72rem; letter-spacing: .16em; text-transform: uppercase;
            color: var(--sun); margin-bottom: .5rem;
        }
        .hero h1, .hero-title {
            font-size: 2.7rem; font-weight: 800; letter-spacing: -.04em;
            color: #fff; line-height: 1.05; margin: 0 0 .5rem 0;
        }
        .hero-sub {
            color: #C9D4DE; font-size: 1.02rem; max-width: 42rem;
            line-height: 1.55; margin-bottom: 1.7rem;
        }
        .hero-stats { display: flex; flex-wrap: wrap; gap: .8rem; }
        .stat {
            background: rgba(255,255,255,.10);
            border: 1px solid rgba(255,255,255,.18);
            backdrop-filter: blur(6px);
            border-radius: 12px;
            padding: .8rem 1.15rem;
            min-width: 8.6rem;
        }
        .stat-label {
            font-size: .68rem; letter-spacing: .1em; text-transform: uppercase;
            color: #A9B7C4; margin-bottom: .28rem;
        }
        .stat-value {
            font-family: var(--mono); font-variant-numeric: tabular-nums;
            font-size: 1.5rem; font-weight: 600; color: #fff; line-height: 1;
        }
        .stat-value.warm { color: var(--sun); }
        .stat-value.cool { color: #4FD1DC; }
        .stat-unit { font-size: .8rem; color: #A9B7C4; margin-left: .18rem; }

        /* ---------- feature cards ---------- */
        .card-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
            gap: 1rem; margin: .4rem 0 1.6rem;
        }
        .fcard {
            background: var(--surface);
            border: 1px solid var(--line);
            border-radius: var(--r);
            padding: 1.35rem 1.3rem;
            box-shadow: var(--shadow);
        }
        .fcard-icon {
            width: 42px; height: 42px; border-radius: 12px;
            display: flex; align-items: center; justify-content: center;
            font-size: 1.25rem; margin-bottom: .85rem;
            background: var(--sun-soft); color: var(--sun-deep);
        }
        .fcard-icon.cool { background: var(--rear-soft); color: var(--rear); }
        .fcard h4 {
            font-size: 1rem; font-weight: 700; margin: 0 0 .3rem;
            color: var(--ink);
        }
        .fcard p {
            font-size: .86rem; color: var(--ink-mid); margin: 0; line-height: 1.5;
        }

        .section-head { margin: 2rem 0 .2rem; }
        .section-head .kicker {
            font-family: var(--mono); font-size: .7rem; letter-spacing: .16em;
            text-transform: uppercase; color: var(--sun-deep);
        }
        .section-head h2 { margin: .2rem 0 .1rem; font-size: 1.6rem; }
        .section-head p { color: var(--ink-mid); font-size: .92rem; margin: 0; }

        /* ---------- streamlit surfaces ---------- */
        section[data-testid="stSidebar"] {
            background: var(--surface); border-right: 1px solid var(--line);
        }
        section[data-testid="stSidebar"] * { font-family: var(--sans); color: var(--ink) !important; }
        section[data-testid="stSidebar"] a[aria-current="page"] {
            background: var(--sun-soft);
            border-left: 3px solid var(--sun-deep);
            border-radius: 0 10px 10px 0; font-weight: 700;
        }

        [data-testid="stMetricValue"] {
            font-family: var(--mono) !important; font-weight: 600 !important;
            font-size: 1.6rem !important; font-variant-numeric: tabular-nums;
            color: var(--ink) !important;
        }
        [data-testid="stMetricLabel"], [data-testid="stMetricLabel"] p {
            font-size: .7rem !important; text-transform: uppercase;
            letter-spacing: .09em; color: var(--ink-mid) !important;
            font-weight: 600 !important;
        }

        div[data-testid="stExpander"],
        div[data-testid="stDataFrame"],
        div[data-testid="stVerticalBlockBorderWrapper"] {
            background: var(--surface) !important;
            border: 1px solid var(--line) !important;
            border-radius: var(--r) !important;
            box-shadow: var(--shadow) !important;
        }
        div[data-testid="stAlert"] {
            border-radius: 12px !important; border: 1px solid var(--line) !important;
        }

        .stButton > button {
            font-family: var(--sans); font-weight: 700; font-size: .88rem;
            border-radius: 10px; padding: .6rem 1.15rem; min-height: 2.7rem;
            box-shadow: none; transition: transform .12s ease, box-shadow .12s ease;
        }
        .stButton > button:hover { transform: translateY(-1px); }
        .stButton > button[kind="primary"],
        .stButton > button[data-testid*="primary"] {
            background: linear-gradient(135deg, var(--sun) 0%, var(--sun-deep) 100%) !important;
            color: #fff !important; border: none !important;
            box-shadow: 0 4px 14px rgba(217,119,6,.28) !important;
        }
        .stButton > button[kind="secondary"],
        .stButton > button[data-testid*="secondary"] {
            background: var(--surface) !important; color: var(--ink) !important;
            border: 1px solid var(--line) !important;
        }
        .stButton > button[kind="secondary"]:hover,
        .stButton > button[data-testid*="secondary"]:hover {
            border-color: var(--sun) !important; color: var(--sun-deep) !important;
        }
        .stButton > button p, .stButton > button div, .stButton > button span { color: inherit !important; }

        /* Option text inside radios and checkboxes is also wrapped in <label>, so the
           widget-label styling below shrank it, upper-cased it and greyed it out
           until the choices rendered as bare unlabelled dots. */
        div[role="radiogroup"] label, div[role="radiogroup"] label p,
        [data-baseweb="radio"] label, [data-baseweb="checkbox"] label,
        [data-testid="stCheckbox"] label, [data-testid="stRadio"] label p {
            color: var(--ink) !important;
            font-size: .9rem !important;
            font-weight: 500 !important;
            text-transform: none !important;
            letter-spacing: normal !important;
        }

        label, [data-testid="stWidgetLabel"], [data-testid="stWidgetLabel"] p {
            color: var(--ink-mid) !important; font-size: .78rem !important;
            font-weight: 600 !important; text-transform: uppercase;
            letter-spacing: .07em; opacity: 1 !important;
        }
        .stTextInput input, .stTextArea textarea, .stNumberInput input,
        input[type="text"], input[type="password"], input[type="number"],
        .stSelectbox div[data-baseweb="select"] {
            background: var(--surface) !important; color: var(--ink) !important;
            -webkit-text-fill-color: var(--ink) !important;
            border: 1px solid var(--line) !important; border-radius: 10px !important;
        }
        .stButton > button:focus-visible, a:focus-visible, input:focus-visible,
        textarea:focus-visible, div[data-baseweb="select"]:focus-within {
            outline: 3px solid var(--sun-deep) !important; outline-offset: 2px !important;
        }
        hr, div[data-testid="stDivider"] { border-top: 1px solid var(--line) !important; }
        [data-testid="stCaptionContainer"] { color: var(--ink-mid) !important; }

        .page-stamp {
            display: inline-block; font-family: var(--mono); font-size: .68rem;
            font-weight: 500; letter-spacing: .14em; text-transform: uppercase;
            color: var(--sun-deep); background: var(--sun-soft);
            border-radius: 999px; padding: 5px 14px; margin-bottom: .6rem;
        }

        /* ---------- solar day bar ---------- */
        .solar-wrap { margin: .1rem 0 1.3rem; }
        .solar-bar {
            position: relative; height: 34px; border-radius: 10px;
            overflow: hidden; background: #16283A;
            box-shadow: inset 0 0 0 1px rgba(255,255,255,.06);
        }
        .solar-day {
            position: absolute; top: 0; bottom: 0;
            background: linear-gradient(90deg, #2F5468 0%, var(--sun) 50%, #2F5468 100%);
        }
        .solar-now { position: absolute; top: -4px; bottom: -4px; width: 2px;
                     background: #fff; box-shadow: 0 0 8px rgba(255,255,255,.8); }
        .solar-scale { display: flex; justify-content: space-between;
                       font-family: var(--mono); font-size: .66rem;
                       color: var(--ink-mid); margin-top: .3rem; }
        .solar-state { font-family: var(--mono); font-size: .72rem;
                       text-transform: uppercase; letter-spacing: .1em; }

        /* ---------- narrow screens ---------- */
        /* Streamlit columns do not stack on a phone, they squeeze: a row of
           five metrics becomes five unreadable slivers. Force them full width
           below the point where that starts to bite. */
        @media (max-width: 640px) {
            /* Streamlit stacks columns itself on narrow screens, one per row.
               For four short readings that wastes half the width and turns
               twenty meters into a hundred rows of scrolling. Overriding it
               needs the full parent > child selector and every box property
               forced, or Streamlit's own rule wins.
               Both test ids are listed because the name changed across
               versions -- "column" on older builds, "stColumn" on newer. */
            div[data-testid="stHorizontalBlock"] {
                flex-wrap: wrap !important;
                gap: 0.5rem !important;
            }
            div[data-testid="stHorizontalBlock"] > div[data-testid="stColumn"],
            div[data-testid="stHorizontalBlock"] > div[data-testid="column"] {
                flex: 0 0 calc(50% - 0.25rem) !important;
                width: calc(50% - 0.25rem) !important;
                min-width: calc(50% - 0.25rem) !important;
                max-width: calc(50% - 0.25rem) !important;
            }

            .hero { padding: 1.6rem 1.2rem 1.4rem; min-height: 0; }
            .hero-title { font-size: 1.8rem; }
            .hero-sub { font-size: .92rem; margin-bottom: 1.1rem; }
            .hero-stats { gap: .5rem; }
            .stat { min-width: calc(50% - .25rem); padding: .6rem .8rem; }
            .stat-value { font-size: 1.2rem; }
            h1 { font-size: 1.9rem !important; }
            .card-grid { grid-template-columns: 1fr; }
            [data-testid="stMetricValue"] { font-size: 1.25rem !important; }
            [data-testid="stMetricLabel"] p { font-size: .62rem !important; }
            .block-container { padding-left: .8rem !important;
                               padding-right: .8rem !important; }
        }

        /* Below about 340px two readings genuinely stop fitting. */
        @media (max-width: 340px) {
            div[data-testid="stHorizontalBlock"] > div[data-testid="stColumn"],
            div[data-testid="stHorizontalBlock"] > div[data-testid="column"] {
                flex: 0 0 100% !important;
                width: 100% !important;
                min-width: 100% !important;
                max-width: 100% !important;
            }
            .stat { min-width: 100%; }
        }

        /* One meter, one card: keeps its four readings visually together
           instead of merging into the next meter's when they wrap. */
        .meter-head {
            font-family: var(--sans); font-weight: 700; font-size: .95rem;
            color: var(--ink); margin: .9rem 0 .2rem;
            border-top: 1px solid var(--line); padding-top: .7rem;
        }
        .meter-head .state { font-weight: 500; font-size: .8rem; margin-left: .5rem; }
        .meter-head .ok { color: var(--power); }
        .meter-head .bad { color: var(--alert); }

        /* The array plan is drawn at a fixed aspect. Shrunk to phone width its
           labels fall to about 6px, so let it keep a legible size and scroll
           sideways instead -- readable and pannable beats present but illegible. */
        .array-scroll { overflow-x: auto; -webkit-overflow-scrolling: touch;
                        border-radius: 16px; box-shadow: var(--shadow); }
        .array-scroll > svg { display: block; min-width: 760px; width: 100%; }
        .array-hint { display: none; font-family: var(--mono); font-size: .7rem;
                      color: var(--ink-mid); margin: .3rem 0 1rem .2rem; }
        @media (max-width: 820px) { .array-hint { display: block; } }

        @media (prefers-reduced-motion: reduce) {
            *, *::before, *::after {
                animation-duration: .001ms !important;
                transition-duration: .001ms !important;
            }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


# Site coordinates, used for the solar day bar. Puchong, Selangor.
SITE_LAT, SITE_LON, SITE_TZ = 3.02, 101.62, 8.0

# datetime.now() returns the server's own clock, which on Streamlit Cloud
# (and most hosts) is UTC, not Malaysia time. Anchoring explicitly to this
# zone keeps "now" correct regardless of what timezone the server runs in.
from zoneinfo import ZoneInfo
SITE_ZONE = ZoneInfo("Asia/Kuala_Lumpur")


def _sun_times(when=None, lat=SITE_LAT, lon=SITE_LON, tz=SITE_TZ):
    """Sunrise, solar noon and sunset for the site, in local hours.

    NOAA's algorithm, worked out here rather than pulled from a library so the
    app has one less dependency to break. Same maths as the gap-filling
    pipeline uses, so the two agree about when daylight is.
    """
    when = when or datetime.now(SITE_ZONE)
    jd = (pd.Timestamp(when.date()) - pd.Timedelta(hours=tz)).to_julian_date() + 0.5
    jc = (jd - 2451545.0) / 36525.0
    L0 = (280.46646 + jc * (36000.76983 + jc * 0.0003032)) % 360.0
    M = 357.52911 + jc * (35999.05029 - 0.0001537 * jc)
    e = 0.016708634 - jc * (0.000042037 + 0.0000001267 * jc)
    Mr = math.radians(M)
    C = (math.sin(Mr) * (1.914602 - jc * (0.004817 + 0.000014 * jc))
         + math.sin(2 * Mr) * (0.019993 - 0.000101 * jc)
         + math.sin(3 * Mr) * 0.000289)
    app = (L0 + C) - 0.00569 - 0.00478 * math.sin(math.radians(125.04 - 1934.136 * jc))
    mo = 23.0 + (26.0 + (21.448 - jc * (46.815 + jc * (0.00059 - jc * 0.001813))) / 60.0) / 60.0
    obl = mo + 0.00256 * math.cos(math.radians(125.04 - 1934.136 * jc))
    decl = math.degrees(math.asin(math.sin(math.radians(obl)) * math.sin(math.radians(app))))
    y = math.tan(math.radians(obl / 2.0)) ** 2
    L0r = math.radians(L0)
    eqt = 4.0 * math.degrees(
        y * math.sin(2 * L0r) - 2 * e * math.sin(Mr)
        + 4 * e * y * math.sin(Mr) * math.cos(2 * L0r)
        - 0.5 * y * y * math.sin(4 * L0r) - 1.25 * e * e * math.sin(2 * Mr))
    noon = (720 - 4 * lon - eqt + tz * 60) / 60.0
    latr, dr = math.radians(lat), math.radians(decl)
    cos_ha = (math.cos(math.radians(90.833)) / (math.cos(latr) * math.cos(dr))
              - math.tan(latr) * math.tan(dr))
    if cos_ha < -1 or cos_ha > 1:      # polar day or night; not here, but safe
        return 0.0, noon, 24.0
    ha = math.degrees(math.acos(cos_ha)) / 15.0
    return noon - ha, noon, noon + ha


def solar_day_bar(when=None):
    """A bar showing where the moment sits in the solar day at the site.

    An empty chart has two very different meanings — the sun is down, or the
    logger has stopped — and nothing on the page distinguished them. This
    does, before you read a single number.
    """
    when = when or datetime.now(SITE_ZONE)
    rise, noon, set_ = _sun_times(when)
    now_h = when.hour + when.minute / 60 + when.second / 3600

    pct = lambda h: max(0.0, min(100.0, 100.0 * h / 24.0))
    left, width, now_x = pct(rise), pct(set_) - pct(rise), pct(now_h)
    daylight = rise <= now_h <= set_

    if daylight:
        state = f"daylight · {(set_ - now_h):.1f} h to sunset"
        colour = "var(--front)"
    else:
        to_rise = (rise - now_h) % 24
        state = f"night · {to_rise:.1f} h to sunrise"
        colour = "var(--ink-mid)"

    fmt = lambda h: f"{int(h):02d}:{int(round((h % 1) * 60)):02d}"
    st.markdown(
        f"""
        <div class="solar-wrap">
          <div class="solar-bar" role="img"
               aria-label="Solar day: sunrise {fmt(rise)}, solar noon
               {fmt(noon)}, sunset {fmt(set_)}. Now {fmt(now_h)}, {state}.">
            <div class="solar-day" style="left:{left}%;width:{width}%"></div>
            <div class="solar-now" style="left:{now_x}%"></div>
          </div>
          <div class="solar-scale">
            <span>00:00</span>
            <span>↑ {fmt(rise)}</span>
            <span>☉ {fmt(noon)}</span>
            <span>↓ {fmt(set_)}</span>
            <span>24:00</span>
          </div>
          <div class="solar-state" style="color:{colour}">{state}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def _bg_image_css(path="BifacialGrid.jpeg"):
    """Inline the array photo as a CSS background.

    Streamlit serves no static folder by default, so a plain url() to a local
    file 404s. Base64 keeps it in the stylesheet and works everywhere,
    including once this is deployed.
    """
    try:
        with open(path, "rb") as fh:
            b64 = base64.b64encode(fh.read()).decode()
        return f"background-image:url('data:image/jpeg;base64,{b64}');"
    except Exception:
        # No photo on disk: fall back to a gradient rather than a blank slab.
        return ("background-image:linear-gradient(120deg,#12283B 0%,"
                "#1E4257 55%,#C77A16 140%);")


def hero(title, subtitle, stats=None, eyebrow="Bifacial PV monitoring"):
    """Full-width banner: array photo, headline, and the numbers that matter.

    `stats` is a list of (label, value, unit, tone) where tone is
    "warm" for front-side / irradiance figures, "cool" for rear-side or
    electrical, and None for neutral.
    """
    chips = ""
    for label, value, unit, tone in (stats or []):
        cls = f"stat-value {tone}" if tone else "stat-value"
        chips += (
            f'<div class="stat"><div class="stat-label">{label}</div>'
            f'<div class="{cls}">{value}'
            f'{f"<span class=stat-unit>{unit}</span>" if unit else ""}</div></div>'
        )
    st.markdown(
        f"""
        <div class="hero" style="{_bg_image_css()}">
          <div class="hero-eyebrow">{eyebrow}</div>
          <div class="hero-title">{title}</div>
          <div class="hero-sub">{subtitle}</div>
          <div class="hero-stats">{chips}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


# Front reference sensors, from the Pi's channel numbering: each row is
# Front A, four rear sensors, Front B.
FRONT_BY_BLOCK = {"B1": (1, 6), "B2": (7, 12), "B3": (13, 18), "B4": (19, 24)}


def rear_sensor_for_meter():
    """Which rear irradiance channel belongs to which meter.

    Each block runs Front A, four rear sensors, Front B, so the nth panel in a
    block pairs with the nth rear sensor. ASSUMED from the numbering rather
    than confirmed against the wiring -- if a per-panel irradiance figure ever
    looks wrong, check this first.
    """
    blocks = [("B1", [11, 12, 13, 14]), ("B2", [15, 16, 17, 18]),
              ("B3", [19, 20, 21, 22]), ("B4", [23, 24, 25, 26])]
    out = {}
    for blk, meters in blocks:
        fa, _fb = FRONT_BY_BLOCK.get(blk, (None, None))
        if fa is None:
            continue
        for i, m in enumerate(meters):
            out[m] = fa + 1 + i
    return out


def latest_per_column(df, cols):
    """Most recent non-null value for each column.

    The logger writes a row every few seconds but does not fill every sensor
    in every row, so df.iloc[-1] is mostly blank -- reading it made eight
    working sensors look like one. Take each column's own last reading.
    """
    out = {}
    if df is None or df.empty:
        return out
    for c in cols:
        if c not in df.columns:
            continue
        ser = pd.to_numeric(df[c], errors="coerce").dropna()
        if len(ser):
            out[c] = float(ser.iloc[-1])
    return out


def split_irradiance(df_live):
    """Average front and rear irradiance separately.

    A single combined figure hides the thing this project measures. Front
    sensors see the direct beam, rear sensors see what the roof reflects, and
    the gap between them IS the bifacial gain -- averaged together they become
    a number that describes neither.

    Returns (front_mean, rear_mean, n_front_live, n_rear_live).
    """
    fronts = [c for pair in FRONT_BY_BLOCK.values() for c in pair]
    if df_live is None or df_live.empty:
        return float("nan"), float("nan"), 0, 0
    latest = latest_per_column(df_live, [f"Irr_{c}" for c in range(1, 25)])
    fv, rv = [], []
    for ch in range(1, 25):
        v = latest.get(f"Irr_{ch}")
        if v is None:
            continue
        (fv if ch in fronts else rv).append(v)
    return (float(np.mean(fv)) if fv else float("nan"),
            float(np.mean(rv)) if rv else float("nan"), len(fv), len(rv))


def array_diagram(values=None, unit="W", title="Live array", front=None,
                  rear=None):
    """A plan view of the roof, each panel lit by what it is producing.

    Four blocks of four bifacial panels, each block bracketed by its two front
    reference sensors. Colour runs from cold slate at zero to bright amber at
    the array's current best, so a dark panel in a bright row is visible
    instantly -- which is the whole question this dashboard exists to answer.

    `values` maps meter id (11-30) to a number. Missing ids render as unlit
    rather than as zero, because "not reporting" and "producing nothing" are
    very different things and should not look the same.
    """
    values = values or {}
    blocks = [("B1", [11, 12, 13, 14]), ("B2", [15, 16, 17, 18]),
              ("B3", [19, 20, 21, 22]), ("B4", [23, 24, 25, 26])]

    # Each block's rear sensors sit between its two front references, so the
    # nth panel pairs with the nth rear sensor. ASSUMED, not confirmed against
    # the physical wiring -- worth checking before trusting a per-panel figure.
    REAR_BY_METER = {}
    for _blk, _meters in blocks:
        _fa, _fb = FRONT_BY_BLOCK.get(_blk, (None, None))
        if _fa is None:
            continue
        for _i, _m in enumerate(_meters):
            REAR_BY_METER[_m] = _fa + 1 + _i

    live = [v for v in values.values() if isinstance(v, (int, float)) and v == v]
    peak = max(live) if live else 0.0
    floor = min(live) if live else 0.0

    # Scale across the spread that is actually present, not from zero.
    #
    # Healthy panels sit within a few percent of each other -- 306 to 322 W in
    # practice -- so a 0-to-peak scale puts every one of them in the top 5% of
    # the ramp and the whole array renders as one flat colour. Stretching the
    # ramp across the observed range makes the differences that matter visible,
    # and the legend states the range so nobody reads brightness as absolute.
    span = peak - floor
    relative = span > 0 and span < peak * 0.35   # tight cluster: stretch it
    lo = floor - span * 0.15 if relative else 0.0

    def tint(v):
        if v is None or v != v:
            return "#243544", "#33475A", 0.28      # unlit: not reporting
        if peak <= 0:
            return "#243544", "#33475A", 0.35
        f = ((v - lo) / (peak - lo)) if peak > lo else 1.0
        f = max(0.0, min(1.0, f))
        # slate -> ember -> amber, so the eye reads intensity as heat
        stops = [(0.0, (36, 53, 68)), (0.45, (124, 74, 30)), (1.0, (242, 160, 61))]
        for i in range(len(stops) - 1):
            a, b = stops[i], stops[i + 1]
            if a[0] <= f <= b[0]:
                t = (f - a[0]) / (b[0] - a[0] or 1)
                rgb = tuple(int(a[1][j] + t * (b[1][j] - a[1][j])) for j in range(3))
                break
        else:
            rgb = stops[-1][1]
        return ("#%02X%02X%02X" % rgb, "#F2A03D" if f > 0.6 else "#5C7086",
                0.55 + 0.45 * f)

    W, H = 980, 386
    pad_x, pad_y = 108, 74      # room for the block label and its sensor ring
    row_h = (H - pad_y - 26) / 4
    panel_w, panel_h = 168, row_h - 16
    parts = [f'<svg viewBox="0 0 {W} {H}" xmlns="http://www.w3.org/2000/svg" '
             f'role="img" aria-label="{title}: 16 bifacial panels in four rows">']
    parts.append('<defs><filter id="glow" x="-40%" y="-40%" width="180%" height="180%">'
                 '<feGaussianBlur stdDeviation="5" result="b"/>'
                 '<feMerge><feMergeNode in="b"/><feMergeNode in="SourceGraphic"/>'
                 '</feMerge></filter></defs>')
    parts.append(f'<rect width="{W}" height="{H}" rx="16" fill="#111E2B"/>')

    for r, (label, ids) in enumerate(blocks):
        y = pad_y + r * row_h
        parts.append(f'<text x="18" y="{y + panel_h/2 + 5}" fill="#8FA0AE" '
                     f'font-family="IBM Plex Mono, monospace" font-size="13" '
                     f'font-weight="600">{label}</text>')
        # front reference sensors sit at each short end of the block
        # Front reference sensors: two per block, one at each short end. They
        # read from the Pi rather than the meters, so they stay hollow until
        # that logger reports -- an empty ring means "no sensor data", which is
        # different from "reading zero" and should not look the same.
        fa, fb = (FRONT_BY_BLOCK.get(label) or (None, None))
        for cx, sid in ((pad_x - 42, fa), (pad_x + 4 * (panel_w + 8) + 14, fb)):
            fv = (front or {}).get(sid)
            has = isinstance(fv, (int, float)) and fv == fv
            fill = "#F2A03D" if has else "none"
            parts.append(f'<circle cx="{cx}" cy="{y + panel_h/2}" r="9" '
                         f'fill="{fill}" fill-opacity="0.85" stroke="#F2A03D" '
                         f'stroke-width="1.6" opacity="{0.95 if has else 0.45}"/>')
            if has:
                parts.append(f'<text x="{cx}" y="{y + panel_h/2 + 23}" '
                             f'text-anchor="middle" fill="#F2A03D" '
                             f'font-family="IBM Plex Mono, monospace" '
                             f'font-size="10">{max(fv, 0):,.0f} W/m&#178;</text>')
            else:
                parts.append(f'<circle cx="{cx}" cy="{y + panel_h/2}" r="2.5" '
                             f'fill="#F2A03D" opacity="0.5"/>')
        for i, dev in enumerate(ids):
            x = pad_x + i * (panel_w + 8)
            v = values.get(dev)
            fill, stroke, op = tint(v)
            glow = ' filter="url(#glow)"' if (v is not None and peak > 0
                                              and v / peak > 0.75) else ''
            parts.append(f'<rect x="{x}" y="{y}" width="{panel_w}" height="{panel_h}" '
                         f'rx="7" fill="{fill}" stroke="{stroke}" stroke-width="1.2" '
                         f'opacity="{op:.2f}"{glow}/>')
            # the cell grid, so it reads as a panel rather than a coloured box
            for k in range(1, 4):
                gx = x + k * panel_w / 4
                parts.append(f'<line x1="{gx}" y1="{y+4}" x2="{gx}" y2="{y+panel_h-4}" '
                             f'stroke="#0C1620" stroke-width="0.8" opacity="0.5"/>')
            parts.append(f'<line x1="{x+4}" y1="{y+panel_h/2}" x2="{x+panel_w-4}" '
                         f'y2="{y+panel_h/2}" stroke="#0C1620" stroke-width="0.8" opacity="0.5"/>')
            if v is None or v != v:
                parts.append(f'<text x="{x + panel_w/2}" y="{y + panel_h/2 + 5}" '
                             f'text-anchor="middle" fill="#8FA0AE" '
                             f'font-family="IBM Plex Mono, monospace" font-size="15" '
                             f'font-weight="600">no reading</text>')
            else:
                parts.append(f'<text x="{x + panel_w/2 - 6}" y="{y + panel_h/2 - 2}" '
                             f'text-anchor="middle" fill="#FFFFFF" '
                             f'font-family="IBM Plex Mono, monospace" font-size="16" '
                             f'font-weight="600" opacity="0.96">{v:,.0f}</text>')
                parts.append(f'<text x="{x + panel_w/2 + 22}" y="{y + panel_h/2 - 2}" '
                             f'fill="#FFFFFF" opacity="0.6" '
                             f'font-family="IBM Plex Mono, monospace" '
                             f'font-size="11">{unit}</text>')

            # Rear irradiance for this panel, underneath its output. Kept in the
            # front-sensor amber so the eye ties the two together, and dimmer
            # than the power figure so the panel still reads power-first.
            rv = (rear or {}).get(dev)
            if isinstance(rv, (int, float)) and rv == rv:
                parts.append(f'<rect x="{x + panel_w/2 - 52}" y="{y + panel_h - 18}" '
                             f'width="104" height="15" rx="3" fill="#0C1620" '
                             f'opacity="0.55"/>')
                parts.append(f'<text x="{x + panel_w/2}" y="{y + panel_h - 7}" '
                             f'text-anchor="middle" fill="#FFD9A0" opacity="0.95" '
                             f'font-family="IBM Plex Mono, monospace" '
                             f'font-size="11" font-weight="500">'
                             f'{max(rv, 0):,.0f} W/m&#178; rear</text>')
            elif rear is not None:
                # Dark text, because this label sits on whatever colour the
                # panel happens to be -- pale grey vanished against a bright
                # amber cell.
                parts.append(f'<rect x="{x + panel_w/2 - 52}" y="{y + panel_h - 18}" '
                             f'width="104" height="15" rx="3" fill="#0C1620" '
                             f'opacity="0.42"/>')
                parts.append(f'<text x="{x + panel_w/2}" y="{y + panel_h - 7}" '
                             f'text-anchor="middle" fill="#DCE4EA" opacity="0.9" '
                             f'font-family="IBM Plex Mono, monospace" '
                             f'font-size="11">no rear sensor</text>')
            parts.append(f'<text x="{x + 8}" y="{y + 15}" fill="#8FA0AE" '
                         f'font-family="IBM Plex Mono, monospace" font-size="10">{dev}</text>')

    parts.append(f'<text x="24" y="26" fill="#E8EEF3" '
                 f'font-family="Manrope, sans-serif" font-size="15" '
                 f'font-weight="700">{title}</text>')

    # A key, because a bare number and an unlabelled circle tell you nothing.
    kx = 24
    parts.append(f'<rect x="{kx}" y="40" width="26" height="13" rx="3" '
                 f'fill="#C87A2A" stroke="#F2A03D" stroke-width="1"/>')
    parts.append(f'<text x="{kx+34}" y="51" fill="#9FB0BE" '
                 f'font-family="IBM Plex Mono, monospace" font-size="11">'
                 f'panel output, watts (rear + front combined)</text>')
    parts.append(f'<circle cx="{kx+352}" cy="46" r="7" fill="#F2A03D" '
                 f'fill-opacity="0.85" stroke="#F2A03D" stroke-width="1.4"/>')
    parts.append(f'<text x="{kx+366}" y="51" fill="#9FB0BE" '
                 f'font-family="IBM Plex Mono, monospace" font-size="11">'
                 f'front reference sensor, W/m&#178;</text>')
    parts.append(f'<circle cx="{kx+640}" cy="46" r="7" fill="none" '
                 f'stroke="#F2A03D" stroke-width="1.4" opacity="0.45"/>')
    parts.append(f'<circle cx="{kx+640}" cy="46" r="2.5" fill="#F2A03D" opacity="0.5"/>')
    parts.append(f'<text x="{kx+654}" y="51" fill="#9FB0BE" '
                 f'font-family="IBM Plex Mono, monospace" font-size="11">'
                 f'sensor not reporting</text>')
    if live:
        rng = (f"{floor:,.0f}-{peak:,.0f} {unit} across the array"
               if span > 0 else f"all at {peak:,.0f} {unit}")
        scale_note = "shading stretched across that range" if relative else \
                     f"shading from 0 to {peak:,.0f} {unit}"
    else:
        rng, scale_note = "no meter readings", "nothing to shade"
    n_front_live = sum(1 for v in (front or {}).values()
                       if isinstance(v, (int, float)) and v == v)
    n_rear_live = sum(1 for v in (rear or {}).values()
                      if isinstance(v, (int, float)) and v == v)
    front_note = (f"{n_front_live} of 8 front and {n_rear_live} of 16 rear "
                  f"sensors reporting")
    parts.append(f'<text x="{pad_x}" y="{H-8}" fill="#6C7C8A" '
                 f'font-family="IBM Plex Mono, monospace" font-size="11">'
                 f'{rng} · {scale_note} · {front_note}</text>')
    parts.append("</svg>")

    st.markdown(
        f'<div class="array-scroll">{"".join(parts)}</div>'
        f'<div class="array-hint">swipe sideways to see the whole array</div>',
        unsafe_allow_html=True,
    )


def section(kicker, heading, blurb=""):
    """A titled section break, so pages read as designed rather than as a
    stack of widgets."""
    st.markdown(
        f"""
        <div class="section-head">
          <div class="kicker">{kicker}</div>
          <h2>{heading}</h2>
          {f"<p>{blurb}</p>" if blurb else ""}
        </div>
        """,
        unsafe_allow_html=True,
    )


def feature_cards(items):
    """Grid of icon cards. `items` is a list of (icon, title, body, tone)."""
    cards = ""
    for icon, title, body, tone in items:
        cls = "fcard-icon cool" if tone == "cool" else "fcard-icon"
        cards += (f'<div class="fcard"><div class="{cls}">{icon}</div>'
                  f'<h4>{title}</h4><p>{body}</p></div>')
    st.markdown(f'<div class="card-grid">{cards}</div>', unsafe_allow_html=True)


def page_stamp(label: str):
    """Renders a small colored pill with the page name at the top of
    a page, in the Solar Admin theme's blue accent."""
    st.markdown(
        f'<div class="page-stamp">{label}</div>',
        unsafe_allow_html=True,
    )


def plot_gauge(value, max_value, title, unit="", color="#C77A16"):
    """A single live reading as an arc.

    Rewritten from a stock Plotly speedometer. The old one banded the track
    red / yellow / green in thirds, which said "low irradiance is bad" -- but
    at 06:00 low irradiance is simply night, and at noon it means cloud.
    Neither is a fault, and colouring it like one is misinformation. The track
    is now neutral and only the reading itself carries colour.
    """
    import plotly.graph_objects as go

    fig = go.Figure(
        go.Indicator(
            mode="gauge+number",
            value=value,
            number={"suffix": f" {unit}",
                    "font": {"size": 34, "family": "IBM Plex Mono, monospace",
                             "color": "#10202B"}},
            gauge={
                "axis": {
                    "range": [0, max_value],
                    "tickcolor": "#D8E0E5",
                    "tickfont": {"size": 10, "family": "IBM Plex Mono, monospace",
                                 "color": "#4A5A66"},
                },
                # a thick arc reading against a hairline track, so the value is
                # the thing you see rather than the decoration behind it
                "bar": {"color": color, "thickness": 0.62},
                "bgcolor": "#EDF1F3",
                "borderwidth": 0,
                "steps": [{"range": [0, max_value], "color": "#EDF1F3"}],
                "threshold": {
                    "line": {"color": "#10202B", "width": 2},
                    "thickness": 0.85,
                    "value": value,
                },
            },
        )
    )
    fig.update_layout(
        height=190,
        margin=dict(l=16, r=16, t=10, b=0),
        paper_bgcolor="rgba(0,0,0,0)",
        font={"family": "IBM Plex Sans, sans-serif", "color": "#10202B"},
    )
    return fig



def plot_line_chart(
    df,
    x_col,
    y_cols,
    x_range=None,
    y_range=None,
    y_title=""
):
    import plotly.graph_objects as go
    import pandas as pd

    fig = go.Figure()

    # Make a copy so we don't modify the original dataframe
    plot_df = df.copy()

    # Make sure timestamp is actually datetime
    plot_df[x_col] = pd.to_datetime(
        plot_df[x_col],
        errors="coerce"
    )

    # Sort chronologically
    plot_df = plot_df.sort_values(x_col)

    for col in y_cols:

        if col not in plot_df.columns:
            continue

        # Convert values to numeric
        y_data = pd.to_numeric(
            plot_df[col],
            errors="coerce"
        )

        # IMPORTANT:
        # Only remove rows where BOTH x and y are invalid.
        trace_df = pd.DataFrame({
            "x": plot_df[x_col],
            "y": y_data,
        }).dropna(subset=["x", "y"])

        # Debugging information
        print(
            f"Plotting {col}: "
            f"{len(trace_df)} valid points"
        )

        fig.add_trace(
            go.Scatter(
                x=trace_df["x"],
                y=trace_df["y"],

                mode="lines",

                name=str(col),

                line=dict(
                    width=2,
                ),

                connectgaps=False,

                hovertemplate=(
                    f"<b>{col}</b><br>"
                    "%{x|%b %d, %Y, %H:%M:%S}<br>"
                    "%{y:.3f}<extra></extra>"
                ),
            )
        )

    fig.update_layout(
        height=380,

        margin=dict(
            l=20,
            r=20,
            t=20,
            b=20,
        ),

        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",

        font={
            "family": "Inter, sans-serif",
            "color": "#1F2937",
        },

        legend={
            "orientation": "h",
            "y": -0.2,
        },

        yaxis_title=y_title,

        showlegend=True,

        hovermode="x unified",
    )

    if x_range is not None:
        fig.update_xaxes(
            range=list(x_range)
        )

    if y_range is not None:
        fig.update_yaxes(
            range=list(y_range)
        )

    return fig
# =========================
# PLOTTING HELPERS
# =========================

def fig_to_image_bytes(fig):
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=300, bbox_inches="tight")
    buf.seek(0)
    return buf


def plot_weather_signals(time, temperatures, irradiances, title="Weather Data", temp_ylim=None, irr_ylim=None):
    # A blank/missing cell in the Time column reads as NaN (a float),
    # which crashes matplotlib's categorical x-axis if it's mixed in
    # with real time strings. pandas' Series.astype(str) doesn't
    # reliably convert every value (some floats can slip through), so
    # convert element-by-element with plain Python str() instead.
    time = [str(t) for t in time]

    fig, ax1 = plt.subplots(figsize=(12, 6))

    for label, temp_values in temperatures.items():
        ax1.plot(time, temp_values, label=str(label))

    ax1.set_xlabel("Time")
    ax1.set_ylabel("Temperature (°C)")
    if temp_ylim is not None:
        ax1.set_ylim(temp_ylim)

    ax2 = ax1.twinx()

    for label, irr_values in irradiances.items():
        ax2.plot(time, irr_values, linestyle="--", label=str(label))

    ax2.set_ylabel("Irradiance (W/m²)")
    if irr_ylim is not None:
        ax2.set_ylim(irr_ylim)

    plt.title(title)

    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, loc="upper left")

    fig.tight_layout()
    return fig


def close_figures(figs):
    """Release matplotlib figures after rendering.

    Streamlit reruns constantly and every rerun built new figures that were
    never closed, so the process leaked memory for as long as the app stayed
    up. Call this after st.pyplot()."""
    for f in (figs.values() if isinstance(figs, dict) else figs):
        try:
            plt.close(f)
        except Exception:
            pass


def plot_irradiance_frequency(df, columns, bin_width=50, max_irr=1200):
    """Returns {column: matplotlib figure}, one histogram per column,
    showing how often that sensor's irradiance readings fell into each
    bin from 0 up to max_irr W/m². Non-numeric/out-of-range values are
    dropped before binning rather than raising."""
    bins = list(range(0, max_irr + bin_width, bin_width))
    figs = {}
    for col in columns:
        values = pd.to_numeric(df[col], errors="coerce").dropna()
        values = values[(values >= 0) & (values <= max_irr)]

        fig, ax = plt.subplots(figsize=(10, 4))
        ax.hist(values, bins=bins, color="#f5a623", edgecolor="black")
        ax.set_xlabel("Irradiance (W/m²)")
        ax.set_ylabel("Frequency")
        ax.set_title(str(col))
        ax.set_xlim(0, max_irr)
        fig.tight_layout()
        figs[col] = fig
    return figs


# =========================
# REPORT DATA BUILDERS & RENDERERS
# =========================

def build_report_data(df, report_title, observation, fig, df_dcm=None):
    """Build uniform report data structure. df_dcm is optional DC
    meter data (Datetime/Device_ID/... columns already standardized
    to created_at/device_id/... by drive_fetch) — when provided, adds
    a per-device DC meter summary table alongside the sensor summary."""
    start_time = f"{df['Date'].iloc[0]} {df['Time'].iloc[0]}" if "Date" in df.columns and "Time" in df.columns else "N/A"
    end_time = f"{df['Date'].iloc[-1]} {df['Time'].iloc[-1]}" if "Date" in df.columns and "Time" in df.columns else "N/A"

    metadata = [
        ("Generated Date", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Total Records", str(df.shape[0])),
        ("Total Columns", str(df.shape[1])),
        ("Start Time", start_time),
        ("End Time", end_time),
        ("Observation Notes", observation if observation else "")
    ]

    columns_list = ", ".join(df.columns)

    numeric_df = df.select_dtypes(include="number")
    numeric_summary = None

    if not numeric_df.empty:
        numeric_summary = []
        for col in numeric_df.columns:
            numeric_summary.append({
                "Column": col,
                "Mean": f"{numeric_df[col].mean():.2f}",
                "Min": f"{numeric_df[col].min():.2f}",
                "Max": f"{numeric_df[col].max():.2f}"
            })

    dcm_summary = None
    if df_dcm is not None and not df_dcm.empty and "device_id" in df_dcm.columns:
        metric_defs = [
            ("voltage_v", "Voltage (V)"),
            ("current_a", "Current (A)"),
            ("active_power_kw", "Power (kW)"),
            ("forward_energy_kwh", "Energy (kWh)"),
        ]
        dcm_summary = []
        for device_id, group in df_dcm.groupby("device_id"):
            device_label = f"Meter {int(device_id)}" if pd.notna(device_id) else "Unknown meter"
            for col, label in metric_defs:
                if col not in group.columns:
                    continue
                vals = pd.to_numeric(group[col], errors="coerce").dropna()
                if vals.empty:
                    continue
                dcm_summary.append({
                    "Column": f"{device_label} — {label}",
                    "Mean": f"{vals.mean():.2f}",
                    "Min": f"{vals.min():.2f}",
                    "Max": f"{vals.max():.2f}",
                })
        if not dcm_summary:
            dcm_summary = None

    return {
        "title": report_title,
        "metadata": metadata,
        "columns": columns_list,
        "numeric_summary": numeric_summary,
        "dcm_summary": dcm_summary,
        "figure": fig
    }


def preview_report_content(df, report_title, observation, fig, df_dcm=None):
    """Display a preview of what will be in the report"""
    report_data = build_report_data(df, report_title, observation, fig, df_dcm=df_dcm)

    # Report Title
    st.markdown(f"# {report_data['title']}")

    # Metadata Table
    metadata_df = pd.DataFrame(report_data['metadata'], columns=["Field", "Value"]) if report_data['metadata'] else pd.DataFrame()
    st.dataframe(metadata_df, use_container_width=True, hide_index=True)

    # Column Overview
    st.markdown("## Column Overview")
    st.write(report_data['columns'])

    # Numeric Summary
    st.markdown("## Numeric Summary")
    if report_data['numeric_summary']:
        summary_df = pd.DataFrame(report_data['numeric_summary'])
        st.dataframe(summary_df, use_container_width=True, hide_index=True)
    else:
        st.info("No numeric columns found")

    # DC Meter Summary
    if report_data['dcm_summary']:
        st.markdown("## DC Meter Summary")
        dcm_df = pd.DataFrame(report_data['dcm_summary'])
        st.dataframe(dcm_df, use_container_width=True, hide_index=True)

    # Weather Graph
    st.markdown("## Weather Graph")
    st.pyplot(report_data['figure'])


def generate_word_report(df, report_title, observation, fig, df_dcm=None):
    report_data = build_report_data(df, report_title, observation, fig, df_dcm=df_dcm)

    doc = Document()
    doc.add_heading(report_data['title'], level=1)

    # Metadata Table
    table = doc.add_table(rows=len(report_data['metadata']), cols=2)
    table.style = "Table Grid"

    for i, (key, value) in enumerate(report_data['metadata']):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = value

    # Column Overview
    doc.add_heading("Column Overview", level=2)
    doc.add_paragraph(report_data['columns'])

    # Numeric Summary
    if report_data['numeric_summary']:
        doc.add_heading("Numeric Summary", level=2)
        summary_table = doc.add_table(rows=len(report_data['numeric_summary']) + 1, cols=4)
        summary_table.style = "Table Grid"

        headers = ["Column", "Mean", "Min", "Max"]
        for col_idx, header in enumerate(headers):
            summary_table.cell(0, col_idx).text = header

        for row_idx, row_data in enumerate(report_data['numeric_summary'], start=1):
            summary_table.cell(row_idx, 0).text = row_data["Column"]
            summary_table.cell(row_idx, 1).text = row_data["Mean"]
            summary_table.cell(row_idx, 2).text = row_data["Min"]
            summary_table.cell(row_idx, 3).text = row_data["Max"]

    # DC Meter Summary
    if report_data['dcm_summary']:
        doc.add_heading("DC Meter Summary", level=2)
        dcm_table = doc.add_table(rows=len(report_data['dcm_summary']) + 1, cols=4)
        dcm_table.style = "Table Grid"

        headers = ["Column", "Mean", "Min", "Max"]
        for col_idx, header in enumerate(headers):
            dcm_table.cell(0, col_idx).text = header

        for row_idx, row_data in enumerate(report_data['dcm_summary'], start=1):
            dcm_table.cell(row_idx, 0).text = row_data["Column"]
            dcm_table.cell(row_idx, 1).text = row_data["Mean"]
            dcm_table.cell(row_idx, 2).text = row_data["Min"]
            dcm_table.cell(row_idx, 3).text = row_data["Max"]

    # Weather Graph
    doc.add_heading("Weather Graph", level=2)
    img_stream = fig_to_image_bytes(report_data['figure'])
    doc.add_picture(img_stream)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_pdf_report(df, report_title, observation, fig, df_dcm=None):
    report_data = build_report_data(df, report_title, observation, fig, df_dcm=df_dcm)

    pdf = FPDF()
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, report_data['title'], new_x="LMARGIN", new_y="NEXT")

    # Metadata Table
    pdf.ln(5)
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(50, 8, "Field", border=1)
    pdf.cell(130, 8, "Value", border=1)
    pdf.ln()

    pdf.set_font("Helvetica", size=10)
    for key, value in report_data['metadata']:
        pdf.cell(50, 8, key, border=1)
        pdf.cell(130, 8, str(value), border=1)
        pdf.ln()

    # Column Overview
    pdf.ln(5)
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Column Overview", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", size=10)
    pdf.multi_cell(0, 8, report_data['columns'])

    # Numeric Summary
    if report_data['numeric_summary']:
        pdf.ln(5)
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 10, "Numeric Summary", new_x="LMARGIN", new_y="NEXT")

        pdf.set_font("Helvetica", "B", 10)
        headers = ["Column", "Mean", "Min", "Max"]
        for header in headers:
            pdf.cell(45, 8, header, border=1)
        pdf.ln()

        pdf.set_font("Helvetica", size=10)
        for row_data in report_data['numeric_summary']:
            pdf.cell(45, 8, str(row_data["Column"]), border=1)
            pdf.cell(45, 8, row_data["Mean"], border=1)
            pdf.cell(45, 8, row_data["Min"], border=1)
            pdf.cell(45, 8, row_data["Max"], border=1)
            pdf.ln()

    # DC Meter Summary
    if report_data['dcm_summary']:
        pdf.ln(5)
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 10, "DC Meter Summary", new_x="LMARGIN", new_y="NEXT")

        pdf.set_font("Helvetica", "B", 10)
        headers = ["Column", "Mean", "Min", "Max"]
        for header in headers:
            pdf.cell(45, 8, header, border=1)
        pdf.ln()

        pdf.set_font("Helvetica", size=10)
        for row_data in report_data['dcm_summary']:
            pdf.cell(45, 8, str(row_data["Column"]), border=1)
            pdf.cell(45, 8, row_data["Mean"], border=1)
            pdf.cell(45, 8, row_data["Min"], border=1)
            pdf.cell(45, 8, row_data["Max"], border=1)
            pdf.ln()

    # Weather Graph
    pdf.ln(5)
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Weather Graph", new_x="LMARGIN", new_y="NEXT")

    img_stream = fig_to_image_bytes(report_data['figure'])

    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
        tmp.write(img_stream.getvalue())
        temp_image_path = tmp.name

    pdf.image(temp_image_path, w=180)

    os.unlink(temp_image_path)

    return bytes(pdf.output())


# =========================
# SUPABASE-DRIVEN SETTINGS & FETCHERS
# =========================

NUM_SENSORS = 24
def get_sensor_logging_config():
    """Get logging mode for each sensor.

    Returns:
        {sensor_id: "normal" | "force_log" | "force_unlog"}
    """
    try:
        res = (
            supabase.table("sensor_logging_config")
            .select("sensor_id, logging_mode")
            .execute()
        )

        return {
            int(row["sensor_id"]): row["logging_mode"]
            for row in (res.data or [])
        }

    except Exception as e:
        print(f"WARNING: couldn't fetch sensor logging config: {e}")
        return {}


def set_sensor_logging_mode(sensor_id, logging_mode):
    """Set one sensor's logging mode."""
    if logging_mode not in ("normal", "force_log", "force_unlog"):
        return False

    try:
        res = (
            supabase.table("sensor_logging_config")
            .upsert(
                {
                    "sensor_id": int(sensor_id),
                    "logging_mode": logging_mode,
                },
                on_conflict="sensor_id",
            )
            .execute()
        )

        return bool(res.data)

    except Exception as e:
        print(
            f"WARNING: couldn't update sensor {sensor_id} "
            f"logging mode: {e}"
        )
        return False
        
def get_forced_sensors():
    try:
        result = (
            supabase
            .table("sensor_logging_config")
            .select("sensor_id")
            .eq("logging_mode", "force_log")
            .execute()
        )
        return [int(row["sensor_id"]) for row in (result.data or [])]
    except Exception as e:
        print(f"WARNING: couldn't fetch forced sensors: {e}")
        return []


def get_excluded_sensors():
    try:
        result = (
            supabase
            .table("sensor_logging_config")
            .select("sensor_id")
            .eq("logging_mode", "force_unlog")
            .execute()
        )
        return [int(row["sensor_id"]) for row in (result.data or [])]
    except Exception as e:
        print(f"WARNING: couldn't fetch excluded sensors: {e}")
        return []

def _safe_query(table: str, limit: int, label: str):
    """Run a Supabase read and surface failures as a message rather than a
    stack trace. Every fetcher used to let an exception escape, so a brief
    connection blip took the whole page down instead of showing stale data."""
    try:
        res = (
            supabase.table(table)
            .select("*")
            .order("created_at", desc=True)
            .limit(limit)
            .execute()
        )
        return res.data or []
    except Exception as exc:
        st.session_state[f"_fetch_error_{table}"] = str(exc)
        st.warning(f"Couldn't load {label} from the database. Showing nothing for now.")
        return []

# Cached like the other two fetchers. Without the decorator this ran a fresh
# database query on every rerun -- including every auto-refresh tick -- and,
# more visibly, `.clear()` does not exist on an uncached function, so the
# Refresh now button raised AttributeError and took the page down.
@st.cache_data(ttl=5)
def fetch_latest_readings(limit=50):
    rows = _safe_query("sensor_readings", limit, "sensor readings")
    if not rows:
        return pd.DataFrame()

    # flatten the jsonb "readings" column into normal columns
    flat_rows = []
    for r in rows:
        flat = {"created_at": r.get("created_at"), "date": r.get("date"), "time": r.get("time")}
        flat.update(r.get("readings") or {})
        flat_rows.append(flat)

    df_live = pd.DataFrame(flat_rows)
    df_live = df_live.sort_values("created_at")  # oldest -> newest for plotting
    return df_live


@st.cache_data(ttl=5)
def fetch_recent_alerts(limit=200):
    rows = _safe_query("sensor_alerts", limit, "sensor alerts")
    if not rows:
        return pd.DataFrame()
    df = pd.DataFrame(rows)
    # errors="coerce": one malformed timestamp used to raise and blank the page
    df["created_at"] = pd.to_datetime(df["created_at"], errors="coerce")
    return df


@st.cache_data(ttl=5)
def fetch_latest_panel_readings(limit=200):
    rows = _safe_query("panel_readings", limit, "panel meter readings")
    if not rows:
        return pd.DataFrame()

    df_panel = pd.DataFrame(rows)
    df_panel = df_panel.sort_values("created_at")  # oldest -> newest for plotting
    return df_panel
