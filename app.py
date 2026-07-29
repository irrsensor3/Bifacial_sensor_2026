import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
from datetime import datetime
import os
import hashlib
import matplotlib.pyplot as plt
from fpdf import FPDF
import tempfile
from supabase import create_client
from streamlit_autorefresh import st_autorefresh

supabase = create_client(
    st.secrets["SUPABASE_URL"],
    st.secrets["SUPABASE_KEY"]
)

# =========================
# AUTHENTICATION
# =========================

SALT = "pv_secure_salt_2026"

def hash_password(password: str) -> str:
    return hashlib.sha256((password + SALT).encode()).hexdigest()

def check_password(password: str) -> bool:
    return hash_password(password) == hash_password("admin123")

def login():
    st.title("🔐 Login")

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Admin Login")
        password = st.text_input("Enter Admin Password", type="password", key="admin_pass")

        if st.button("Login as Admin"):
            if check_password(password):
                st.session_state.auth = True
                st.session_state.user_role = "admin"
                st.success("Admin access granted")
                st.rerun()
            else:
                st.error("Wrong password")

    with col2:
        st.subheader("Guest Access")
        if st.button("Login as Guest"):
            st.session_state.auth = True
            st.session_state.user_role = "guest"
            st.success("Guest access granted")
            st.rerun()

# =========================
# PLOTTING FUNCTION
# =========================

def fig_to_image_bytes(fig):
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=300, bbox_inches="tight")
    buf.seek(0)
    return buf

def plot_weather_signals(time, temperatures, irradiances, title="Weather Data"):
    fig, ax1 = plt.subplots(figsize=(12, 6))

    for label, temp_values in temperatures.items():
        ax1.plot(time, temp_values, label=label)

    ax1.set_xlabel("Time")
    ax1.set_ylabel("Temperature (°C)")

    ax2 = ax1.twinx()

    for label, irr_values in irradiances.items():
        ax2.plot(time, irr_values, linestyle="--", label=label)

    ax2.set_ylabel("Irradiance (W/m²)")

    plt.title(title)

    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, loc="upper left")

    fig.tight_layout()
    return fig

# =========================
# REPORT DATA BUILDER
# =========================

def build_report_data(df, report_title, observation, fig):
    """Build uniform report data structure"""
    start_time = f"{df['Date'].iloc[0]} {df['Time'].iloc[0]}" if "Date" in df.columns and "Time" in df.columns else "N/A"
    end_time = f"{df['Date'].iloc[-1]} {df['Time'].iloc[-1]}" if "Date" in df.columns and "Time" in df.columns else "N/A"

    metadata = [
        ("Generated Date", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Total Records", str(df.shape[0])),
        ("Total Columns", str(df.shape[1])),
        ("Start Time", start_time),
        ("End Time", end_time),
        ("Observation Notes", observation if observation else "")
    ]

    columns_list = ", ".join(df.columns)

    numeric_df = df.select_dtypes(include="number")
    numeric_summary = None

    if not numeric_df.empty:
        numeric_summary = []
        for col in numeric_df.columns:
            numeric_summary.append({
                "Column": col,
                "Mean": f"{numeric_df[col].mean():.2f}",
                "Min": f"{numeric_df[col].min():.2f}",
                "Max": f"{numeric_df[col].max():.2f}"
            })

    return {
        "title": report_title,
        "metadata": metadata,
        "columns": columns_list,
        "numeric_summary": numeric_summary,
        "figure": fig
    }

# =========================
# PREVIEW FUNCTION
# =========================

def preview_report_content(df, report_title, observation, fig):
    """Display a preview of what will be in the report"""
    report_data = build_report_data(df, report_title, observation, fig)

    # Report Title
    st.markdown(f"# {report_data['title']}")

    # Metadata Table
    metadata_df = pd.DataFrame(report_data['metadata'], columns=["Field", "Value"])
    st.dataframe(metadata_df, use_container_width=True, hide_index=True)

    # Column Overview
    st.markdown("## Column Overview")
    st.write(report_data['columns'])

    # Numeric Summary
    st.markdown("## Numeric Summary")
    if report_data['numeric_summary']:
        summary_df = pd.DataFrame(report_data['numeric_summary'])
        st.dataframe(summary_df, use_container_width=True, hide_index=True)
    else:
        st.info("No numeric columns found")

    # Weather Graph
    st.markdown("## Weather Graph")
    st.pyplot(report_data['figure'])

# =========================
# WORD REPORT
# =========================

def generate_word_report(df, report_title, observation, fig):
    report_data = build_report_data(df, report_title, observation, fig)

    doc = Document()
    doc.add_heading(report_data['title'], level=1)

    # Metadata Table
    table = doc.add_table(rows=len(report_data['metadata']), cols=2)
    table.style = "Table Grid"

    for i, (key, value) in enumerate(report_data['metadata']):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = value

    # Column Overview
    doc.add_heading("Column Overview", level=2)
    doc.add_paragraph(report_data['columns'])

    # Numeric Summary
    if report_data['numeric_summary']:
        doc.add_heading("Numeric Summary", level=2)
        summary_table = doc.add_table(rows=len(report_data['numeric_summary']) + 1, cols=4)
        summary_table.style = "Table Grid"

        headers = ["Column", "Mean", "Min", "Max"]
        for col_idx, header in enumerate(headers):
            summary_table.cell(0, col_idx).text = header

        for row_idx, row_data in enumerate(report_data['numeric_summary'], start=1):
            summary_table.cell(row_idx, 0).text = row_data["Column"]
            summary_table.cell(row_idx, 1).text = row_data["Mean"]
            summary_table.cell(row_idx, 2).text = row_data["Min"]
            summary_table.cell(row_idx, 3).text = row_data["Max"]

    # Weather Graph
    doc.add_heading("Weather Graph", level=2)
    img_stream = fig_to_image_bytes(report_data['figure'])
    doc.add_picture(img_stream)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =========================
# PDF REPORT
# =========================

def generate_pdf_report(df, report_title, observation, fig):
    report_data = build_report_data(df, report_title, observation, fig)

    pdf = FPDF()
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, report_data['title'], new_x="LMARGIN", new_y="NEXT")

    # Metadata Table
    pdf.ln(5)
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(50, 8, "Field", border=1)
    pdf.cell(130, 8, "Value", border=1)
    pdf.ln()

    pdf.set_font("Helvetica", size=10)
    for key, value in report_data['metadata']:
        pdf.cell(50, 8, key, border=1)
        pdf.cell(130, 8, str(value), border=1)
        pdf.ln()

    # Column Overview
    pdf.ln(5)
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Column Overview", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", size=10)
    pdf.multi_cell(0, 8, report_data['columns'])

    # Numeric Summary
    if report_data['numeric_summary']:
        pdf.ln(5)
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 10, "Numeric Summary", new_x="LMARGIN", new_y="NEXT")

        pdf.set_font("Helvetica", "B", 10)
        headers = ["Column", "Mean", "Min", "Max"]
        for header in headers:
            pdf.cell(45, 8, header, border=1)
        pdf.ln()

        pdf.set_font("Helvetica", size=10)
        for row_data in report_data['numeric_summary']:
            pdf.cell(45, 8, str(row_data["Column"]), border=1)
            pdf.cell(45, 8, row_data["Mean"], border=1)
            pdf.cell(45, 8, row_data["Min"], border=1)
            pdf.cell(45, 8, row_data["Max"], border=1)
            pdf.ln()

    # Weather Graph
    pdf.ln(5)
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Weather Graph", new_x="LMARGIN", new_y="NEXT")

    img_stream = fig_to_image_bytes(report_data['figure'])

    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
        tmp.write(img_stream.getvalue())
        temp_image_path = tmp.name

    pdf.image(temp_image_path, w=180)

    os.unlink(temp_image_path)

    return bytes(pdf.output())

# =========================
# REMOTE FORCE-LOG SETTING (pi_settings table, id=1)
# =========================

def get_force_log_status():
    """Best-effort read of the current force_log_below_zero flag.
    Returns False (normal/safe behavior) if the table/row doesn't
    exist yet or the request fails, so a Supabase hiccup never shows
    a false 'currently forcing' state."""
    try:
        res = (
            supabase.table("pi_settings")
            .select("force_log_below_zero")
            .eq("id", 1)
            .execute()
        )
        if res.data:
            return bool(res.data[0].get("force_log_below_zero", False))
    except Exception:
        pass
    return False


def set_force_log_status(value: bool):
    """Best-effort write of the force_log_below_zero flag. Returns
    True on success so the caller can show an error if it didn't
    actually go through."""
    try:
        supabase.table("pi_settings").update(
            {"force_log_below_zero": value}
        ).eq("id", 1).execute()
        return True
    except Exception:
        return False

# =========================
# MAIN APP
# =========================

if "auth" not in st.session_state:
    st.session_state.auth = False
    st.session_state.user_role = None

if not st.session_state.auth:
    login()
    st.stop()

st.sidebar.subheader("Session")
st.sidebar.write(f"Role: {st.session_state.user_role.capitalize()}")

if st.sidebar.button("🚪 Logout"):
    st.session_state.auth = False
    st.session_state.user_role = None
    st.rerun()

st.title("📊 Bifacial PV Data Logging System")

file = st.file_uploader("Upload CSV", type=["csv"])

report_title = st.text_input("Report Title", "Bifacial PV Performance Report")
observation = st.text_area("Observation Notes")

if file is not None:

    df = pd.read_csv(file)

    st.subheader("📊 Data Preview")
    st.dataframe(df.head(100))

    st.subheader("📌 Dataset Info")
    st.write(f"Rows: {df.shape[0]}")
    st.write(f"Columns: {df.shape[1]}")

    time = df["Time"] if "Time" in df.columns else df.index

    numeric_cols = df.select_dtypes(include="number").columns.tolist()

    st.subheader("📈 Graph Configuration")

    selected_temps = st.multiselect(
        "Select Temperature Columns",
        numeric_cols,
        default=[c for c in numeric_cols if "temp" in c.lower()]
    )

    selected_irradiance = st.multiselect(
        "Select Irradiance Columns",
        numeric_cols,
        default=[c for c in numeric_cols if "irr" in c.lower()]
    )

    temperatures = {col: df[col].tolist() for col in selected_temps}
    irradiances = {col: df[col].tolist() for col in selected_irradiance}

    if selected_temps or selected_irradiance:
        fig = plot_weather_signals(time, temperatures, irradiances)
        st.pyplot(fig)

        st.subheader("📄 Generate Reports")

        if st.button("👁️ Preview Report"):
            st.session_state.show_preview = True

        if st.session_state.get("show_preview"):
            with st.expander("📋 Report Preview", expanded=True):
                preview_report_content(df, report_title, observation, fig)

            st.divider()
            col1, col2 = st.columns(2)

            with col1:
                report = generate_word_report(df, report_title, observation, fig)
                st.download_button(
                    label="⬇️ Download Word Report",
                    data=report,
                    file_name="PV_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            with col2:
                report = generate_pdf_report(df, report_title, observation, fig)
                st.download_button(
                    label="⬇️ Download PDF Report",
                    data=report,
                    file_name="PV_Report.pdf",
                    mime="application/pdf"
                )

    if st.button("Test Supabase"):
        supabase.table("pi_commands").update({"command": "hello"}).eq("id", 1).execute()
        st.success("Database updated!")

if st.session_state.user_role == "admin":
    st.divider()
    st.subheader("🔴 Admin Controls")

    if st.button("🔄 Reboot Raspberry Pi"):

        supabase.table("pi_commands") \
            .update({"command": "reboot"}) \
            .eq("id", 1) \
            .execute()

        st.success("Reboot command sent.")

    if st.button("⚫ Shutdown Raspberry Pi"):

        supabase.table("pi_commands") \
            .update({"command": "shutdown"}) \
            .eq("id", 1) \
            .execute()
        st.success("Shutdown command sent.")

    # -------------------------------------------------
    # FORCE LOGGING BELOW 0°C
    # -------------------------------------------------
    st.divider()
    st.subheader("🌡️ Force Logging Below 0°C")

    force_log_on = get_force_log_status()

    status_col, button_col = st.columns([1, 2])

    with status_col:
        if force_log_on:
            st.markdown(
                "<div style='display:flex;align-items:center;gap:8px;'>"
                "<div style='width:16px;height:16px;border-radius:50%;"
                "background:#2ecc71;box-shadow:0 0 8px #2ecc71;'></div>"
                "<span><b>FORCE LOG: ON</b></span></div>",
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                "<div style='display:flex;align-items:center;gap:8px;'>"
                "<div style='width:16px;height:16px;border-radius:50%;"
                "background:#e74c3c;box-shadow:0 0 8px #e74c3c;'></div>"
                "<span><b>FORCE LOG: OFF</b></span></div>",
                unsafe_allow_html=True,
            )

    with button_col:
        if force_log_on:
            if st.button("⏹️ Stop Force Logging (resume sub-zero cutoff)"):
                if set_force_log_status(False):
                    st.success("Force logging disabled — sub-zero readings will be discarded again.")
                    st.rerun()
                else:
                    st.error("Couldn't reach Supabase — try again.")
        else:
            if st.button("▶️ Force Logging (ignore sub-zero cutoff)"):
                if set_force_log_status(True):
                    st.success("Force logging enabled — the Pi will log sub-zero readings as-is.")
                    st.rerun()
                else:
                    st.error("Couldn't reach Supabase — try again.")

    st.caption(
        "When ON, the Pi keeps logging every sensor's temperature even if it "
        "reads below 0°C, instead of discarding it as invalid. The Pi only "
        "checks this once a minute, so it can take up to ~60s to take effect."
    )
else:
    st.divider()
    st.info("ℹ️ Admin controls are not available in guest mode.")

# =========================
# LIVE SENSOR DATA
# =========================

st.divider()
st.subheader("Live Sensor Data")

refresh_seconds = st.slider("Auto-refresh interval (seconds)", 5, 60, 15)
st_autorefresh(interval=refresh_seconds * 1000, key="live_refresh")

@st.cache_data(ttl=5)
def fetch_latest_readings(limit=50):
    response = (
        supabase.table("sensor_readings")
        .select("*")
        .order("created_at", desc=True)
        .limit(limit)
        .execute()
    )
    rows = response.data
    if not rows:
        return pd.DataFrame()

    # flatten the jsonb "readings" column into normal columns
    flat_rows = []
    for r in rows:
        flat = {"created_at": r["created_at"], "date": r["date"], "time": r["time"]}
        flat.update(r["readings"] or {})
        flat_rows.append(flat)

    df_live = pd.DataFrame(flat_rows)
    df_live = df_live.sort_values("created_at")  # oldest -> newest for plotting
    return df_live

df_live = fetch_latest_readings()

if df_live.empty:
    st.info("No live data yet — waiting for the Pi to push a sample.")
else:
    latest = df_live.iloc[-1]
    st.caption(f"Last update: {latest['date']} {latest['time']}")

    irr_cols = [c for c in df_live.columns if c.startswith("Irr_")]
    temp_cols = [c for c in df_live.columns if c.startswith("Temp_")]

    # quick "latest values" snapshot
    cols = st.columns(4)
    for i, col in enumerate(irr_cols[:4]):
        val = latest[col]
        cols[i].metric(col, f"{val:.1f} W/m²" if val is not None else "—")

    # live trend chart — reuse your existing plotting logic if you prefer
    selected_live_irr = st.multiselect(
        "Irradiance sensors to plot (live)", irr_cols, default=irr_cols[:1]
    )
    if selected_live_irr:
        st.line_chart(df_live.set_index("created_at")[selected_live_irr])

    with st.expander("Raw live data table"):
        st.dataframe(df_live, use_container_width=True, hide_index=True)

    # -------------------------------------------------
    # SENSORS CURRENTLY BELOW 0°C (not logging)
    # -------------------------------------------------
    st.markdown("### 🌡️ Sensors Below 0°C")

    @st.cache_data(ttl=5)
    def fetch_recent_alerts(limit=200):
        response = (
            supabase.table("sensor_alerts")
            .select("*")
            .order("created_at", desc=True)
            .limit(limit)
            .execute()
        )
        rows = response.data
        if not rows:
            return pd.DataFrame()
        df = pd.DataFrame(rows)
        df["created_at"] = pd.to_datetime(df["created_at"])
        return df

    df_alerts = fetch_recent_alerts()

    if df_alerts.empty:
        st.success("No sub-zero alerts recorded — all sensors logging normally.")
    else:
        # a sensor's most recent alert — if it's recent (within ~2.5x the
        # Pi's once-a-minute check), treat it as still currently invalid
        latest_per_sensor = df_alerts.sort_values("created_at").groupby("sensor_id").tail(1)
        now_utc = pd.Timestamp.now(tz="UTC")
        currently_invalid = latest_per_sensor[
            now_utc - latest_per_sensor["created_at"] < pd.Timedelta(seconds=150)
        ]

        if currently_invalid.empty:
            st.success("No sensors currently below 0°C.")
        else:
            st.error(f"{len(currently_invalid)} sensor(s) currently below 0°C and not logging:")
            for _, row in currently_invalid.sort_values("sensor_id").iterrows():
                st.markdown(
                    f"🔴 **Sensor {int(row['sensor_id'])}** — {row['temp_c']}°C "
                    f"@ {row['created_at'].strftime('%H:%M:%S')} "
                    f"(bus {row['bus']}, addr {row['address']})"
                )

        with st.expander("Recent alert history"):
            st.dataframe(df_alerts, use_container_width=True, hide_index=True)

# =========================
# LIVE PANEL METER DATA (DCM3366, RS485 bus, up to 8 meters)
# =========================

st.divider()
st.subheader("Live Panel Meter Data")

@st.cache_data(ttl=5)
def fetch_latest_panel_readings(limit=200):
    response = (
        supabase.table("panel_readings")
        .select("*")
        .order("created_at", desc=True)
        .limit(limit)
        .execute()
    )
    rows = response.data
    if not rows:
        return pd.DataFrame()

    df_panel = pd.DataFrame(rows)
    df_panel = df_panel.sort_values("created_at")  # oldest -> newest for plotting
    return df_panel

df_panel = fetch_latest_panel_readings()

if df_panel.empty:
    st.info("No live panel meter data yet — waiting for the mini PC to push a sample.")
else:
    # most recent reading per meter (device_id)
    latest_per_device = (
        df_panel.sort_values("created_at").groupby("device_id").tail(1).sort_values("device_id")
    )
    latest_overall = df_panel.iloc[-1]
    st.caption(f"Last update: {latest_overall['created_at']}")

    for _, row in latest_per_device.iterrows():
        device_label = f"Meter {int(row['device_id'])}"
        has_error = row.get("error") not in (None, "No error")

        status_dot = "🔴" if has_error else "🟢"
        cols = st.columns(5)
        cols[0].markdown(f"**{status_dot} {device_label}**")
        cols[1].metric("Voltage", f"{row['voltage_v']:.1f} V" if row["voltage_v"] is not None else "—")
        cols[2].metric("Current", f"{row['current_a']:.2f} A" if row["current_a"] is not None else "—")
        cols[3].metric("Power", f"{row['active_power_kw']:.3f} kW" if row["active_power_kw"] is not None else "—")
        cols[4].metric("Energy", f"{row['forward_energy_kwh']:.1f} kWh" if row["forward_energy_kwh"] is not None else "—")

        if has_error:
            st.caption(f"⚠️ {device_label}: {row.get('error')}")

    st.markdown("### Trend")
    device_ids = sorted(df_panel["device_id"].dropna().unique().tolist())
    selected_devices = st.multiselect(
        "Meters to plot", device_ids, default=device_ids[:1], format_func=lambda d: f"Meter {int(d)}"
    )
    metric_choice = st.radio(
        "Metric", ["voltage_v", "current_a", "active_power_kw"],
        format_func=lambda m: {"voltage_v": "Voltage (V)", "current_a": "Current (A)", "active_power_kw": "Power (kW)"}[m],
        horizontal=True,
    )

    if selected_devices:
        pivot = df_panel[df_panel["device_id"].isin(selected_devices)].pivot_table(
            index="created_at", columns="device_id", values=metric_choice
        )
        pivot.columns = [f"Meter {int(c)}" for c in pivot.columns]
        st.line_chart(pivot)

    with st.expander("Raw panel meter data table"):
        st.dataframe(df_panel, use_container_width=True, hide_index=True)
